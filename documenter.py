import os
import zipfile
import tempfile
import re
import argparse
# from openai import OpenAI
from docx import Document
from docx.shared import Pt
import cohere
# ======== Settings ========
UPLOADS_DIR = "uploads"
OUTPUTS_DIR = "outputs"
MAX_TOKENS = 1024
MODEL_NAME = "command-r"

# ======== Initialize OpenAI Client ========


# ======== Initialize Cohere Client ========
api_key = os.getenv("COHERE_API_KEY")
if not api_key:
    raise ValueError("COHERE_API_KEY environment variable not set.")

client = cohere.Client(api_key)


# ======== Helper Functions ========

def extract_zip(zip_path, extract_to):
    with zipfile.ZipFile(zip_path, 'r') as zip_ref:
        zip_ref.extractall(extract_to)

def read_files(root_dir):
    file_data = []
    for root, dirs, files in os.walk(root_dir):
        for file in files:
            if file.endswith(('.py', '.js', '.java', '.cpp', '.c', '.html', '.css', '.ts', '.php', '.go')):
                file_path = os.path.join(root, file)
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                    file_data.append((file, content))
                except Exception as e:
                    print(f"[!] Error reading {file_path}: {e}")
    return file_data

def generate_text(prompt):
    try:
        response = client.chat(
            model="command-r",
            message=prompt,
            temperature=0.5,
            max_tokens=MAX_TOKENS
        )
        return response.text.strip()
    except Exception as e:
        print(f"[!] API Error: {e}")
        return "(Error during generation)"


def generate_project_summary(all_code_combined):
    prompt = f"""
You are a project documentation expert.
Given the full codebase provided below, write a Project Level Summary covering these points:

- Purpose
- Intuition
- Technologies Used
- Scope of Improvement

Important Rules:
- Only answer based on the actual code provided.
- No assumptions, no invented tech stacks.
- Keep the answers crisp, technical, and clean.

Codebase:
text
{all_code_combined[:6000]}

"""
    return generate_text(prompt)

def generate_file_summary(file_name, file_content):
    prompt = f"""
You are a professional code documentation assistant.
Analyze the following file carefully.

Return exactly 3 sections, each beginning with the marker:
- Purpose:
- Working:
- Errors:

Example:
Purpose: ...
Working: ...
Errors: ...

Be clear, technical and concise. No assumptions.

File Name: {file_name}
Code:
text
{file_content[:2000]}

"""
    return generate_text(prompt)

def parse_sections(text):
    match = re.findall(r"(Purpose|Working|Errors):\s*(.*?)(?=(?:Purpose|Working|Errors):|$)", text, re.DOTALL)
    section_dict = {}
    for key, val in match:
        clean_val = val.strip().replace('\n', ' ')
        if clean_val.lower().startswith(key.lower()):
            clean_val = clean_val[len(key):].lstrip(':').strip()
        section_dict[key] = clean_val if clean_val else "(Not provided)"
    for key in ["Purpose", "Working", "Errors"]:
        if key not in section_dict:
            section_dict[key] = "(Not provided)"
    return section_dict

def create_docx(project_summary, file_summaries, output_path):
    doc = Document()
    doc.add_heading('Project Documentation', 0)
    doc.add_heading('Project Level Summary', level=1)
    para = doc.add_paragraph(project_summary)
    para.style.font.size = Pt(11)
    doc.add_heading('File Level Summaries', level=1)
    for file_name, summary in file_summaries.items():
        doc.add_heading(file_name, level=2)
        para = doc.add_paragraph(summary)
        para.style.font.size = Pt(10)
    doc.save(output_path)

def create_markdown(project_summary, file_summaries, output_path):
    lines = []
    lines.append("# Project Documentation\n")
    lines.append("## Project Level Summary\n")
    lines.append(project_summary.strip())
    lines.append("\n---\n")
    lines.append("## File Level Summaries\n")
    for file_name, summary in file_summaries.items():
        lines.append(f"### {file_name}\n")
        parsed = parse_sections(summary)
        for section in ["Purpose", "Working", "Errors"]:
            lines.append(f"- **{section}:** {parsed.get(section, '(Not provided)')}")
        lines.append("\n---\n")
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))

# ======== Main Pipeline ========

def main(zip_filename):
    zip_path = os.path.join(UPLOADS_DIR, zip_filename)
    base_name = os.path.splitext(zip_filename)[0]
    output_docx_path = os.path.join(OUTPUTS_DIR, f"{base_name}_documentation.docx")
    output_md_path = os.path.join(OUTPUTS_DIR, f"{base_name}_documentation.md")

    with tempfile.TemporaryDirectory() as tempdir:
        print("[+] Extracting ZIP...")
        extract_zip(zip_path, tempdir)

        print("[+] Reading files...")
        file_data = read_files(tempdir)

        if not file_data:
            print("[!] No code files found.")
            return

        all_code_combined = "\n\n".join(content for _, content in file_data)

        print("[+] Generating project level summary...")
        project_summary = generate_project_summary(all_code_combined)

        print("[+] Generating file level summaries...")
        file_summaries = {}
        for file_name, file_content in file_data:
            summary = generate_file_summary(file_name, file_content)
            file_summaries[file_name] = summary

        print("[+] Creating DOCX...")
        os.makedirs(OUTPUTS_DIR, exist_ok=True)
        create_docx(project_summary, file_summaries, output_docx_path)

        print("[+] Creating Markdown...")
        create_markdown(project_summary, file_summaries, output_md_path)

        print(f"[✓] Documentation created successfully at:\n  - {output_docx_path}\n  - {output_md_path}")

# ======== Entry Point ========

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate documentation from a ZIP file of source code.")
    parser.add_argument("zipfile", type=str, help="Name of the ZIP file inside the uploads/ directory")
    args = parser.parse_args()

    zip_file_path = os.path.join(UPLOADS_DIR, args.zipfile)
    if not os.path.isfile(zip_file_path):
        print(f"[!] File not found: {zip_file_path}")
    else:
        main(args.zipfile)
